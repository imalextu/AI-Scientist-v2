#!/usr/bin/env python3
from __future__ import annotations

import argparse
from datetime import datetime
from io import BytesIO
import json
import queue
import re
import threading
import uuid
from pathlib import Path
from typing import Any, Dict

from flask import (
    Flask,
    Response,
    jsonify,
    request,
    send_file,
    send_from_directory,
    stream_with_context,
)

from paper_agent.config import load_config_from_text
from paper_agent.workflow import ThesisWorkflow, WorkflowCancelledError

SCRIPT_DIR = Path(__file__).resolve().parent
UI_DIR = SCRIPT_DIR / "web_ui"

app = Flask(__name__, static_folder=str(UI_DIR), static_url_path="/ui")

DEFAULTS: Dict[str, str] = {
    "config_path": "",
    "topic_path": "",
    "config_text": "",
    "topic_text": "",
}
JOBS: Dict[str, Dict[str, Any]] = {}
RUNNING_JOBS: Dict[str, Dict[str, Any]] = {}
JOBS_LOCK = threading.Lock()

STAGE_SEQUENCE = ["literature", "idea", "review", "outline", "paper"]
STAGE_ORDER = {name: idx + 1 for idx, name in enumerate(STAGE_SEQUENCE)}
STAGE_FILE_SPECS: Dict[str, tuple[str, str]] = {
    "literature": ("00_literature.json", "json"),
    "idea": ("01_idea.json", "json"),
    "review": ("00_literature_review.md", "text"),
    "outline": ("02_outline.json", "json"),
    "paper": ("03_thesis.md", "text"),
}
AUX_OUTPUT_FILE_SPECS: Dict[str, tuple[str, str]] = {
    "paper_initial": ("03a_thesis_initial.md", "text"),
    "paper_audit": ("03b_thesis_audit.md", "text"),
}
CACHE_ROOT = SCRIPT_DIR / ".cache_store"
CACHE_ID_PATTERN = re.compile(r"^[a-zA-Z0-9_-]+$")
INVALID_FILENAME_CHARS_RE = re.compile(r'[\\/:*?"<>|]+')
MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
MARKDOWN_ORDERED_LIST_RE = re.compile(r"^(\s*)\d+[.)]\s+(.+)$")
MARKDOWN_UNORDERED_LIST_RE = re.compile(r"^(\s*)[-*+]\s+(.+)$")
MARKDOWN_QUOTE_RE = re.compile(r"^>\s?(.*)$")
MARKDOWN_HORIZONTAL_RULE_RE = re.compile(r"^([-*_])\1{2,}$")


def parse_stage(value: Any) -> str | None:
    stage = str(value or "").strip().lower()
    if stage in STAGE_ORDER:
        return stage
    return None


def stages_up_to(stage: str) -> list[str]:
    return [name for name in STAGE_SEQUENCE if STAGE_ORDER[name] <= STAGE_ORDER[stage]]


def resumable_stages_for_cache(stage: str) -> list[str]:
    stage_idx = STAGE_SEQUENCE.index(stage)
    max_idx = min(stage_idx + 1, len(STAGE_SEQUENCE) - 1)
    return STAGE_SEQUENCE[: max_idx + 1]


def normalize_cache_id(raw_value: Any) -> str:
    cache_id = str(raw_value or "").strip()
    if not cache_id or not CACHE_ID_PATTERN.fullmatch(cache_id):
        raise ValueError("cache_id 非法")
    return cache_id


def ensure_cache_root() -> Path:
    CACHE_ROOT.mkdir(parents=True, exist_ok=True)
    return CACHE_ROOT


def resolve_cache_dir(cache_id: str) -> Path:
    return ensure_cache_root() / cache_id


def write_json_file(path: Path, payload: Any) -> None:
    path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def read_json_file(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def read_cache_manifest(cache_id: str) -> Dict[str, Any]:
    manifest_path = resolve_cache_dir(cache_id) / "cache_manifest.json"
    if not manifest_path.exists():
        raise FileNotFoundError(f"未找到缓存: {cache_id}")
    payload = read_json_file(manifest_path)
    if not isinstance(payload, dict):
        raise ValueError(f"缓存元数据格式错误: {cache_id}")
    payload["cache_id"] = cache_id
    return payload


def read_cache_outputs(cache_id: str) -> Dict[str, str]:
    cache_dir = resolve_cache_dir(cache_id)
    outputs: Dict[str, str] = {}
    for stage in STAGE_SEQUENCE:
        file_name, content_type = STAGE_FILE_SPECS[stage]
        path = cache_dir / file_name
        if not path.exists():
            continue
        raw_text = path.read_text(encoding="utf-8")
        if content_type == "json":
            parsed = json.loads(raw_text)
            outputs[stage] = json.dumps(parsed, ensure_ascii=False, indent=2)
        else:
            outputs[stage] = raw_text
    for stage, (file_name, content_type) in AUX_OUTPUT_FILE_SPECS.items():
        path = cache_dir / file_name
        if not path.exists():
            continue
        raw_text = path.read_text(encoding="utf-8")
        if content_type == "json":
            parsed = json.loads(raw_text)
            outputs[stage] = json.dumps(parsed, ensure_ascii=False, indent=2)
        else:
            outputs[stage] = raw_text
    return outputs


def build_resume_cache_payload(cache_id: str) -> Dict[str, Any]:
    cache_dir = resolve_cache_dir(cache_id)
    payload: Dict[str, Any] = {}

    literature_file = cache_dir / STAGE_FILE_SPECS["literature"][0]
    if literature_file.exists():
        literature_data = read_json_file(literature_file)
        if not isinstance(literature_data, list):
            raise ValueError("缓存中的 00_literature.json 不是数组")
        payload["literature_items"] = literature_data

    idea_file = cache_dir / STAGE_FILE_SPECS["idea"][0]
    if idea_file.exists():
        idea_data = read_json_file(idea_file)
        if not isinstance(idea_data, dict):
            raise ValueError("缓存中的 01_idea.json 不是对象")
        payload["idea_data"] = idea_data

    review_file = cache_dir / STAGE_FILE_SPECS["review"][0]
    if review_file.exists():
        review_text = review_file.read_text(encoding="utf-8")
        payload["review_text"] = review_text

    outline_file = cache_dir / STAGE_FILE_SPECS["outline"][0]
    if outline_file.exists():
        outline_data = read_json_file(outline_file)
        if not isinstance(outline_data, dict):
            raise ValueError("缓存中的 02_outline.json 不是对象")
        payload["outline_data"] = outline_data

    return payload


def normalize_manifest(payload: Any, fallback_cache_id: str) -> Dict[str, Any] | None:
    if not isinstance(payload, dict):
        return None
    stage = parse_stage(payload.get("stage"))
    if not stage:
        return None
    cache_id = str(payload.get("cache_id") or fallback_cache_id).strip() or fallback_cache_id
    created_at = str(payload.get("created_at") or "")
    stages_included = payload.get("stages_included")
    if not isinstance(stages_included, list) or not stages_included:
        stages_included = stages_up_to(stage)
    else:
        cleaned: list[str] = []
        for item in stages_included:
            parsed = parse_stage(item)
            if parsed and parsed not in cleaned:
                cleaned.append(parsed)
        if not cleaned:
            cleaned = stages_up_to(stage)
        stages_included = cleaned
    resumable = payload.get("resumable_stages")
    if not isinstance(resumable, list) or not resumable:
        resumable = resumable_stages_for_cache(stage)
    else:
        cleaned_resume: list[str] = []
        for item in resumable:
            parsed = parse_stage(item)
            if parsed and parsed not in cleaned_resume:
                cleaned_resume.append(parsed)
        if not cleaned_resume:
            cleaned_resume = resumable_stages_for_cache(stage)
        resumable = cleaned_resume
    return {
        "cache_id": cache_id,
        "created_at": created_at,
        "stage": stage,
        "stages_included": stages_included,
        "resumable_stages": resumable,
        "title": str(payload.get("title") or "").strip(),
        "topic_preview": str(payload.get("topic_preview") or "").strip(),
    }


def resolve_input_path(path_str: str) -> Path:
    candidate = Path(path_str).expanduser()
    if candidate.is_absolute():
        return candidate
    if candidate.exists():
        return candidate.resolve()
    return (SCRIPT_DIR / candidate).resolve()


def read_optional_text(path: Path) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8")


def set_defaults(config_path: Path, topic_path: Path) -> None:
    DEFAULTS["config_path"] = str(config_path)
    DEFAULTS["topic_path"] = str(topic_path)
    DEFAULTS["config_text"] = read_optional_text(config_path)
    DEFAULTS["topic_text"] = read_optional_text(topic_path)


def format_sse_event(event: str, payload: Dict[str, Any]) -> str:
    content = json.dumps(payload, ensure_ascii=False)
    return f"event: {event}\ndata: {content}\n\n"


def sanitize_docx_filename(raw_name: Any) -> str:
    candidate = str(raw_name or "").strip()
    if not candidate:
        return "03_thesis.docx"
    candidate = INVALID_FILENAME_CHARS_RE.sub("_", candidate)
    candidate = re.sub(r"\s+", " ", candidate).strip(" .")
    if not candidate:
        candidate = "03_thesis"
    if not candidate.lower().endswith(".docx"):
        candidate = f"{candidate}.docx"
    return candidate


def normalize_markdown_inline_text(raw_text: str) -> str:
    text = str(raw_text or "")
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"_([^_]+)_", r"\1", text)
    text = re.sub(r"~~([^~]+)~~", r"\1", text)
    return text.strip()


def flush_paragraph_buffer(document: Any, lines: list[str]) -> None:
    if not lines:
        return
    paragraph_text = " ".join(
        normalize_markdown_inline_text(line) for line in lines if str(line).strip()
    ).strip()
    lines.clear()
    if paragraph_text:
        document.add_paragraph(paragraph_text)


def add_code_block(document: Any, lines: list[str], pt_factory: Any) -> None:
    if not lines:
        return
    code_text = "\n".join(lines).rstrip("\n")
    lines.clear()
    if not code_text:
        return
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = pt_factory(10)


def add_list_item(
    document: Any,
    content: str,
    ordered: bool,
    pt_factory: Any,
    indent_level: int = 0,
) -> None:
    text = normalize_markdown_inline_text(content)
    if not text:
        return
    style_name = "List Number" if ordered else "List Bullet"
    try:
        paragraph = document.add_paragraph(text, style=style_name)
    except KeyError:
        marker = "1." if ordered else "•"
        paragraph = document.add_paragraph(f"{marker} {text}")
    if indent_level > 0:
        paragraph.paragraph_format.left_indent = pt_factory(18 * indent_level)


def markdown_to_docx_buffer(markdown_text: str) -> BytesIO:
    try:
        from docx import Document
        from docx.shared import Pt
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "未安装 python-docx，请先执行：pip install -r requirements.txt"
        ) from exc

    document = Document()
    lines = str(markdown_text or "").replace("\r\n", "\n").replace("\r", "\n").split("\n")
    paragraph_lines: list[str] = []
    code_lines: list[str] = []
    in_code_block = False

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                add_code_block(document, code_lines, Pt)
                in_code_block = False
            else:
                flush_paragraph_buffer(document, paragraph_lines)
                in_code_block = True
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            flush_paragraph_buffer(document, paragraph_lines)
            continue

        heading_match = MARKDOWN_HEADING_RE.match(stripped)
        if heading_match:
            flush_paragraph_buffer(document, paragraph_lines)
            level = min(len(heading_match.group(1)), 6)
            heading_text = normalize_markdown_inline_text(heading_match.group(2))
            if heading_text:
                document.add_heading(heading_text, level=level)
            continue

        ordered_match = MARKDOWN_ORDERED_LIST_RE.match(line)
        unordered_match = MARKDOWN_UNORDERED_LIST_RE.match(line)
        if ordered_match or unordered_match:
            flush_paragraph_buffer(document, paragraph_lines)
            match = ordered_match if ordered_match else unordered_match
            if match:
                indent_spaces = len(match.group(1).expandtabs(2))
                add_list_item(
                    document,
                    match.group(2),
                    ordered=bool(ordered_match),
                    pt_factory=Pt,
                    indent_level=min(indent_spaces // 2, 8),
                )
            continue

        quote_match = MARKDOWN_QUOTE_RE.match(stripped)
        if quote_match:
            flush_paragraph_buffer(document, paragraph_lines)
            quote_text = normalize_markdown_inline_text(quote_match.group(1))
            if quote_text:
                paragraph = document.add_paragraph(quote_text)
                paragraph.paragraph_format.left_indent = Pt(18)
            continue

        if MARKDOWN_HORIZONTAL_RULE_RE.match(stripped):
            flush_paragraph_buffer(document, paragraph_lines)
            document.add_paragraph("-" * 40)
            continue

        paragraph_lines.append(stripped)

    if in_code_block:
        add_code_block(document, code_lines, Pt)
    flush_paragraph_buffer(document, paragraph_lines)

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output


@app.get("/")
def index() -> Response:
    return send_from_directory(UI_DIR, "index.html")


@app.get("/api/initial")
def initial() -> Response:
    return jsonify(DEFAULTS)


@app.post("/api/export/word")
def export_word() -> Response:
    payload = request.get_json(silent=True) or {}
    markdown_text = str(payload.get("markdown", ""))
    if not markdown_text.strip():
        return jsonify({"error": "markdown 不能为空"}), 400

    file_name = sanitize_docx_filename(payload.get("file_name"))
    try:
        docx_buffer = markdown_to_docx_buffer(markdown_text)
    except Exception as exc:
        return jsonify({"error": f"导出 Word 失败: {exc}"}), 500

    return send_file(
        docx_buffer,
        as_attachment=True,
        download_name=file_name,
        mimetype=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )


@app.post("/api/cache/save")
def save_cache() -> Response:
    payload = request.get_json(silent=True) or {}
    stage = parse_stage(payload.get("stage"))
    if not stage:
        return jsonify({"error": "stage 非法"}), 400

    config_text = str(payload.get("config_text", "")).strip()
    topic_text = str(payload.get("topic_text", "")).strip()
    outputs = payload.get("outputs")

    if not config_text:
        return jsonify({"error": "config_text 不能为空"}), 400
    if not topic_text:
        return jsonify({"error": "topic_text 不能为空"}), 400
    if not isinstance(outputs, dict):
        return jsonify({"error": "outputs 格式错误"}), 400

    required_stages = stages_up_to(stage)
    included_stages: list[str] = []
    parsed_outputs: Dict[str, Any] = {}
    for stage_name in required_stages:
        raw_text = str(outputs.get(stage_name, "")).strip()
        if not raw_text:
            return jsonify({"error": f"{stage_name} 阶段内容为空，无法缓存"}), 400

        file_name, content_type = STAGE_FILE_SPECS[stage_name]
        if content_type == "json":
            try:
                parsed_outputs[stage_name] = json.loads(raw_text)
            except json.JSONDecodeError:
                return (
                    jsonify(
                        {"error": f"{file_name} 不是合法 JSON，当前阶段可能尚未完成"}
                    ),
                    400,
                )
        else:
            parsed_outputs[stage_name] = raw_text
        included_stages.append(stage_name)

    cache_id = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"
    cache_dir = resolve_cache_dir(cache_id)
    cache_dir.mkdir(parents=True, exist_ok=False)

    (cache_dir / "config.yaml").write_text(f"{config_text}\n", encoding="utf-8")
    (cache_dir / "00_topic.md").write_text(f"{topic_text}\n", encoding="utf-8")

    for stage_name in included_stages:
        file_name, content_type = STAGE_FILE_SPECS[stage_name]
        target_file = cache_dir / file_name
        if content_type == "json":
            write_json_file(target_file, parsed_outputs[stage_name])
        else:
            target_file.write_text(parsed_outputs[stage_name], encoding="utf-8")

    if stage == "paper":
        for aux_stage, (file_name, _content_type) in AUX_OUTPUT_FILE_SPECS.items():
            aux_text = str(outputs.get(aux_stage, "")).strip()
            if not aux_text:
                continue
            (cache_dir / file_name).write_text(aux_text, encoding="utf-8")

    title = ""
    idea_obj = parsed_outputs.get("idea")
    if isinstance(idea_obj, dict):
        title = str(idea_obj.get("thesis_title_cn") or "").strip()
    topic_preview = re.sub(r"\s+", " ", topic_text).strip()[:120]

    manifest = {
        "cache_id": cache_id,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "stage": stage,
        "stages_included": included_stages,
        "resumable_stages": resumable_stages_for_cache(stage),
        "title": title,
        "topic_preview": topic_preview,
    }
    write_json_file(cache_dir / "cache_manifest.json", manifest)
    return jsonify({"cache": manifest})


@app.get("/api/cache/list")
def list_caches() -> Response:
    root = ensure_cache_root()
    caches: list[Dict[str, Any]] = []
    for path in root.iterdir():
        if not path.is_dir():
            continue
        manifest_path = path / "cache_manifest.json"
        if not manifest_path.exists():
            continue
        try:
            raw = read_json_file(manifest_path)
            manifest = normalize_manifest(raw, path.name)
            if not manifest:
                continue
            caches.append(manifest)
        except Exception:
            continue
    caches.sort(key=lambda item: str(item.get("created_at") or ""), reverse=True)
    return jsonify({"caches": caches})


@app.get("/api/cache/<cache_id>")
def get_cache(cache_id: str) -> Response:
    try:
        normalized_cache_id = normalize_cache_id(cache_id)
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400

    try:
        raw_manifest = read_cache_manifest(normalized_cache_id)
        manifest = normalize_manifest(raw_manifest, normalized_cache_id)
        if not manifest:
            raise ValueError("缓存元数据缺少必要字段")
        outputs = read_cache_outputs(normalized_cache_id)
    except FileNotFoundError:
        return jsonify({"error": "缓存不存在"}), 404
    except Exception as exc:
        return jsonify({"error": f"读取缓存失败: {exc}"}), 500

    cache_dir = resolve_cache_dir(normalized_cache_id)
    config_text = read_optional_text(cache_dir / "config.yaml")
    topic_text = read_optional_text(cache_dir / "00_topic.md")
    return jsonify(
        {
            "cache": manifest,
            "config_text": config_text,
            "topic_text": topic_text,
            "outputs": outputs,
        }
    )


@app.post("/api/jobs")
def create_job() -> Response:
    payload = request.get_json(silent=True) or {}
    config_text = str(payload.get("config_text", "")).strip()
    topic_text = str(payload.get("topic_text", "")).strip()
    forced_title_raw = str(payload.get("title", "")).strip()
    forced_title = forced_title_raw or None
    resume_cache_id_raw = str(payload.get("resume_cache_id", "")).strip()
    resume_from_stage = ""
    resume_to_stage = ""

    if not config_text:
        return jsonify({"error": "config_text 不能为空"}), 400
    if not topic_text:
        return jsonify({"error": "topic_text 不能为空"}), 400

    if resume_cache_id_raw:
        try:
            resume_cache_id = normalize_cache_id(resume_cache_id_raw)
        except ValueError as exc:
            return jsonify({"error": str(exc)}), 400
        try:
            raw_manifest = read_cache_manifest(resume_cache_id)
            manifest = normalize_manifest(raw_manifest, resume_cache_id)
            if not manifest:
                return jsonify({"error": "缓存元数据格式错误"}), 400
        except FileNotFoundError:
            return jsonify({"error": "缓存不存在"}), 404
        except Exception as exc:
            return jsonify({"error": f"读取缓存失败: {exc}"}), 400

        parsed_resume_from = parse_stage(payload.get("resume_from_stage"))
        if not parsed_resume_from:
            return jsonify({"error": "resume_from_stage 非法"}), 400
        resume_from_stage = parsed_resume_from

        resume_to_raw = str(payload.get("resume_to_stage", "")).strip()
        if resume_to_raw:
            parsed_resume_to = parse_stage(resume_to_raw)
            if not parsed_resume_to:
                return jsonify({"error": "resume_to_stage 非法"}), 400
            resume_to_stage = parsed_resume_to

        if not resume_to_stage:
            resume_to_stage = "paper"

        if STAGE_ORDER[resume_from_stage] > STAGE_ORDER[resume_to_stage]:
            return jsonify({"error": "resume_from_stage 不能晚于 resume_to_stage"}), 400

        resumable_stages = manifest.get("resumable_stages")
        if isinstance(resumable_stages, list) and resumable_stages:
            if resume_from_stage not in resumable_stages:
                return (
                    jsonify(
                        {
                            "error": (
                                f"该缓存不可从 {resume_from_stage} 继续，"
                                f"可选: {', '.join(resumable_stages)}"
                            )
                        }
                    ),
                    400,
                )
    else:
        resume_cache_id = ""

    job_id = uuid.uuid4().hex
    with JOBS_LOCK:
        JOBS[job_id] = {
            "config_text": config_text,
            "topic_text": topic_text,
            "forced_title": forced_title,
            "resume_cache_id": resume_cache_id,
            "resume_from_stage": resume_from_stage,
            "resume_to_stage": resume_to_stage,
        }
    return jsonify({"job_id": job_id})


@app.post("/api/jobs/<job_id>/cancel")
def cancel_job(job_id: str) -> Response:
    with JOBS_LOCK:
        pending_job = JOBS.pop(job_id, None)
        running_job = RUNNING_JOBS.get(job_id)

    if pending_job is not None:
        return jsonify({"status": "cancelled", "message": "任务已取消（尚未启动）"})

    if not running_job:
        return jsonify({"error": "任务不存在或已结束"}), 404

    cancel_event = running_job.get("cancel_event")
    if isinstance(cancel_event, threading.Event):
        cancel_event.set()

    return jsonify({"status": "cancelling", "message": "已发送停止请求，正在终止任务"})


@app.get("/api/jobs/<job_id>/events")
def stream_job(job_id: str) -> Response:
    with JOBS_LOCK:
        job = JOBS.pop(job_id, None)
    if not job:
        return jsonify({"error": "任务不存在或已开始执行"}), 404

    cancel_event = threading.Event()
    with JOBS_LOCK:
        RUNNING_JOBS[job_id] = {"cancel_event": cancel_event}

    def event_stream() -> Any:
        updates: queue.Queue[tuple[str, Dict[str, Any]] | None] = queue.Queue(
            maxsize=2048
        )
        stream_closed = threading.Event()

        def emit(event: str, payload: Dict[str, Any]) -> None:
            if stream_closed.is_set():
                return
            try:
                updates.put_nowait((event, payload))
            except queue.Full:
                try:
                    updates.get_nowait()
                except queue.Empty:
                    return
                try:
                    updates.put_nowait((event, payload))
                except queue.Full:
                    return

        def worker() -> None:
            try:
                config = load_config_from_text(
                    job["config_text"], project_root=SCRIPT_DIR
                )
                resume_cache_payload: Dict[str, Any] | None = None
                resume_from_stage: str | None = None
                resume_cache_id = str(job.get("resume_cache_id") or "").strip()
                resume_to_stage = str(job.get("resume_to_stage") or "").strip()
                if resume_cache_id:
                    resume_cache_payload = build_resume_cache_payload(resume_cache_id)
                    resume_from_stage = str(job.get("resume_from_stage") or "").strip()
                    if resume_to_stage:
                        config.workflow.run_until_stage = resume_to_stage
                    emit(
                        "resume_started",
                        {
                            "cache_id": resume_cache_id,
                            "resume_from_stage": resume_from_stage,
                            "resume_to_stage": config.workflow.run_until_stage,
                        },
                    )
                workflow = ThesisWorkflow(
                    config,
                    event_handler=emit,
                    cancel_checker=cancel_event.is_set,
                )
                run_dir = workflow.run(
                    topic_text=job["topic_text"],
                    forced_title=job["forced_title"],
                    resume_from_stage=resume_from_stage,
                    resume_cache=resume_cache_payload,
                )
                if cancel_event.is_set():
                    emit("job_cancelled", {"message": "任务已停止"})
                else:
                    emit("done", {"run_dir": str(run_dir)})
            except WorkflowCancelledError:
                emit("job_cancelled", {"message": "任务已停止"})
            except Exception as exc:
                emit("job_error", {"message": str(exc)})
            finally:
                with JOBS_LOCK:
                    RUNNING_JOBS.pop(job_id, None)
                try:
                    updates.put_nowait(None)
                except queue.Full:
                    return

        threading.Thread(target=worker, daemon=True).start()
        yield format_sse_event("status", {"message": "任务已启动"})

        try:
            while True:
                try:
                    item = updates.get(timeout=15)
                except queue.Empty:
                    yield ": keep-alive\n\n"
                    continue

                if item is None:
                    break
                event, payload = item
                yield format_sse_event(event, payload)
        finally:
            stream_closed.set()

    headers = {
        "Cache-Control": "no-cache",
        "X-Accel-Buffering": "no",
    }
    return Response(
        stream_with_context(event_stream()),
        mimetype="text/event-stream",
        headers=headers,
    )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="管理学论文生成器本地 Web 界面")
    parser.add_argument(
        "--config",
        type=str,
        default="config.yaml",
        help="默认加载的配置文件路径",
    )
    parser.add_argument(
        "--topic-file",
        type=str,
        default="examples/topic_example.md",
        help="默认加载的题目描述文件路径",
    )
    parser.add_argument("--host", type=str, default="127.0.0.1", help="监听地址")
    parser.add_argument("--port", type=int, default=7860, help="监听端口")
    parser.add_argument("--debug", action="store_true", help="开启调试模式")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    config_path = resolve_input_path(args.config)
    topic_path = resolve_input_path(args.topic_file)
    set_defaults(config_path, topic_path)

    if not UI_DIR.exists():
        raise FileNotFoundError(f"未找到前端目录: {UI_DIR}")

    app.run(host=args.host, port=args.port, debug=args.debug, threaded=True)


if __name__ == "__main__":
    main()
